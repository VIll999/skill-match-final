"""
Text Extraction Service
Handles extraction of text from various document formats (PDF, DOCX, TXT)
"""
import logging
import io
import tempfile
import os
from pathlib import Path
from typing import Optional, Dict, Any
import re

# PDF processing
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# DOCX processing
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# Fallback for other formats
try:
    import textract
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Exception raised when text extraction fails"""
    pass


class TextExtractor:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['application/pdf', '.pdf'],
            'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx'],
            'doc': ['application/msword', '.doc'],
            'txt': ['text/plain', '.txt'],
            'rtf': ['application/rtf', '.rtf'],
            'odt': ['application/vnd.oasis.opendocument.text', '.odt']
        }
        
        # Log available extractors
        logger.info(f"TextExtractor initialized with:")
        logger.info(f"  PDFPlumber: {HAS_PDFPLUMBER}")
        logger.info(f"  PyMuPDF: {HAS_PYMUPDF}")
        logger.info(f"  python-docx: {HAS_PYTHON_DOCX}")
        logger.info(f"  textract: {HAS_TEXTRACT}")
    
    def extract_text(self, file_content: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """
        Extract text from file content
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type (optional)
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Determine file type
            file_ext = Path(filename).suffix.lower()
            detected_format = self._detect_format(file_ext, content_type)
            
            logger.info(f"Extracting text from {filename} (format: {detected_format})")
            
            # Extract text based on format
            if detected_format == 'pdf':
                text = self._extract_pdf(file_content)
            elif detected_format == 'docx':
                text = self._extract_docx(file_content)
            elif detected_format == 'txt':
                text = self._extract_txt(file_content)
            else:
                # Fallback to textract
                text = self._extract_with_textract(file_content, filename)
            
            # Normalize whitespace
            normalized_text = self._normalize_text(text)
            
            # Extract metadata
            metadata = self._extract_metadata(normalized_text, filename, detected_format)
            
            return {
                'text': normalized_text,
                'raw_text': text,
                'filename': filename,
                'format': detected_format,
                'metadata': metadata,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return {
                'text': '',
                'raw_text': '',
                'filename': filename,
                'format': 'unknown',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    def _detect_format(self, file_ext: str, content_type: str = None) -> str:
        """Detect file format from extension and content type"""
        
        # Check by extension first
        for format_name, (mime_types, extensions) in self.supported_formats.items():
            if file_ext in extensions:
                return format_name
        
        # Check by content type
        if content_type:
            for format_name, (mime_types, extensions) in self.supported_formats.items():
                if content_type in mime_types:
                    return format_name
        
        # Default to textract fallback
        return 'unknown'
    
    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber or PyMuPDF"""
        text = ""
        
        # Try pdfplumber first (better for tables and layout)
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                if text.strip():
                    logger.debug(f"PDF extracted with pdfplumber: {len(text)} characters")
                    return text
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Fallback to PyMuPDF (faster for simple text)
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                
                logger.debug(f"PDF extracted with PyMuPDF: {len(text)} characters")
                return text
                
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        raise TextExtractionError("No PDF extraction library available")
    
    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        if not HAS_PYTHON_DOCX:
            raise TextExtractionError("python-docx not available")
        
        try:
            doc = Document(io.BytesIO(file_content))
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            result = '\n'.join(text)
            logger.debug(f"DOCX extracted: {len(result)} characters")
            return result
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise TextExtractionError(f"DOCX extraction failed: {e}")
    
    def _extract_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8')
            logger.debug(f"TXT extracted (UTF-8): {len(text)} characters")
            return text
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                text = file_content.decode('latin-1')
                logger.debug(f"TXT extracted (latin-1): {len(text)} characters")
                return text
            except UnicodeDecodeError:
                try:
                    # Final fallback with error handling
                    text = file_content.decode('utf-8', errors='replace')
                    logger.debug(f"TXT extracted (UTF-8 with errors): {len(text)} characters")
                    return text
                except Exception as e:
                    raise TextExtractionError(f"Text file decoding failed: {e}")
    
    def _extract_with_textract(self, file_content: bytes, filename: str) -> str:
        """Extract text using textract as fallback"""
        if not HAS_TEXTRACT:
            raise TextExtractionError("textract not available")
        
        try:
            # textract requires a file path, so we need to save temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Extract text
                text = textract.process(tmp_file_path).decode('utf-8')
                logger.debug(f"Textract extracted: {len(text)} characters")
                return text
            finally:
                # Clean up temp file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Textract extraction failed: {e}")
            raise TextExtractionError(f"Textract extraction failed: {e}")
    
    def _normalize_text(self, text: str) -> str:
        """Normalize whitespace and clean text"""
        if not text:
            return ""
        
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Clean up common artifacts
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _extract_metadata(self, text: str, filename: str, format_type: str) -> Dict[str, Any]:
        """Extract metadata from text"""
        metadata = {
            'char_count': len(text),
            'word_count': len(text.split()) if text else 0,
            'line_count': len(text.splitlines()) if text else 0,
            'format': format_type,
            'filename': filename
        }
        
        # Try to extract basic resume sections
        if text:
            sections = self._detect_resume_sections(text)
            metadata['sections'] = sections
            
            # Extract contact information
            contact_info = self._extract_contact_info(text)
            metadata['contact_info'] = contact_info
        
        return metadata
    
    def _detect_resume_sections(self, text: str) -> Dict[str, bool]:
        """Detect common resume sections"""
        sections = {}
        text_lower = text.lower()
        
        # Common section headers
        section_patterns = {
            'experience': r'\b(experience|work\s+history|employment|professional\s+experience)\b',
            'education': r'\b(education|academic|qualifications|degrees?)\b',
            'skills': r'\b(skills|technical\s+skills|competencies|abilities)\b',
            'projects': r'\b(projects|portfolio|work\s+samples)\b',
            'certifications': r'\b(certifications?|certificates?|licenses?)\b',
            'awards': r'\b(awards?|honors?|achievements?|recognition)\b',
            'publications': r'\b(publications?|papers?|articles?|research)\b',
            'languages': r'\b(languages?|linguistic)\b'
        }
        
        for section, pattern in section_patterns.items():
            sections[section] = bool(re.search(pattern, text_lower))
        
        return sections
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract basic contact information"""
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern (basic)
        phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?[2-9]\d{2}\)?[-.\s]?[2-9]\d{2}[-.\s]?\d{4}\b'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9_-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # GitHub pattern
        github_pattern = r'github\.com/[A-Za-z0-9_-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = github_match.group()
        
        return contact_info
    
    def get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Get information about supported formats"""
        formats = {}
        
        for format_name, (mime_types, extensions) in self.supported_formats.items():
            available = True
            
            # Check if required libraries are available
            if format_name == 'pdf':
                available = HAS_PDFPLUMBER or HAS_PYMUPDF
            elif format_name == 'docx':
                available = HAS_PYTHON_DOCX
            elif format_name in ['doc', 'rtf', 'odt']:
                available = HAS_TEXTRACT
            
            formats[format_name] = {
                'mime_types': mime_types,
                'extensions': extensions,
                'available': available
            }
        
        return formats
    
    def validate_file(self, filename: str, content_type: str = None, file_size: int = None) -> Dict[str, Any]:
        """Validate if file can be processed"""
        file_ext = Path(filename).suffix.lower()
        
        # Check if format is supported
        supported = False
        for format_name, (mime_types, extensions) in self.supported_formats.items():
            if file_ext in extensions or (content_type and content_type in mime_types):
                supported = True
                break
        
        # Check file size (10MB limit)
        max_size = 10 * 1024 * 1024  # 10MB
        size_ok = file_size is None or file_size <= max_size
        
        return {
            'supported': supported,
            'size_ok': size_ok,
            'max_size': max_size,
            'detected_format': self._detect_format(file_ext, content_type)
        }